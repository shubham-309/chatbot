from typing import List, Optional
from pydantic import BaseModel, Field
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from openpyxl import load_workbook
import PyPDF2
import streamlit_authenticator as stauth
import yaml
from yaml.loader import SafeLoader
from langchain.schema import Document
import streamlit as st
import os
from pinecone import Pinecone
from langchain_pinecone import PineconeVectorStore
from langchain_openai import OpenAIEmbeddings
from docx import Document as DocxDocument  # Import for DOCX handling
from dotenv import load_dotenv

load_dotenv()

api_key = os.getenv("PINECONE_API_KEY")

st.set_page_config(layout="wide", page_title="FZcompare Admin")

st.markdown(
    """
    <style>
    .block-container {
        padding-top: 0rem; /* Reduce top padding */
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Load config file with error handling
try:
    with open('config.yaml', 'r') as file:
        config = yaml.load(file, Loader=SafeLoader)
except FileNotFoundError:
    st.error("config.yaml not found. Please ensure it exists in the working directory.")
    st.stop()
except yaml.YAMLError as e:
    st.error(f"Error parsing config.yaml: {e}")
    st.stop()

# Initialize authenticator
authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)

# Define the structured output model
class CompanyInfo(BaseModel):
    name: str = Field(..., description="Name of the Provider")
    package: str = Field(description="Name of the Package")
    number_of_visas: Optional[int] = Field(None, description="Number of visas")
    number_of_shareholders: Optional[int] = Field(None, description="Number of shareholders")
    office_required: Optional[bool] = Field(None, description="Is an office required?")
    cost: Optional[float] = Field(None, description="Associated cost")
    activity: Optional[str] = Field(description="Activities Allowed in the freezone, e.g. consultancy, media, trading, freelance, and industrial activities etc")

# Define a wrapper model for multiple CompanyInfo instances
class CompanyInfoList(BaseModel):
    companies: List[CompanyInfo]

# Initialize OpenAI model and embeddings
llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.0)
embeddings = OpenAIEmbeddings()

# Create a Pinecone client instance with your API key
pc = Pinecone()
index = os.getenv('PINECONE_INDEX_NAME')

# Initialize the parser for CompanyInfoList
parser = PydanticOutputParser(pydantic_object=CompanyInfoList)

# Define the prompt template for structured extraction
prompt = PromptTemplate(
    template="Extract structured information from the following text:\n\n"
             "- Name of the Provider\n"
             "- Name of the Package\n"
             "- Number of visas\n"
             "- Number of shareholders\n"
             "- Is an office required? (True/False)\n"
             "- Associated cost\n\n"
             "- Activities Allowed in the freezone, e.g. consultancy, media, trading, freelance, and industrial activities etc\n\n"
             "Provide the output in JSON format.\n\n"
             "Text: {text}\n\n"
             "{format_instructions}",
    input_variables=["text"],
    partial_variables={"format_instructions": parser.get_format_instructions()},
)

enhancement_prompt = PromptTemplate(
    template="create meaningful sentences from this text each sentence should mention:\n\n"
             "- Name of the Provider\n"
             "- Name of the Package\n"
             "- Number of visas\n"
             "- Number of shareholders\n"
             "- Is an office required? (True/False)\n"
             "- Associated cost\n\n"
             "- Activities Allowed in the freezone, e.g. consultancy, media, trading, freelance, and industrial activities etc\n\n"
             "Text: {text}\n\n",
    input_variables=["text"]
)

text_enhance_prompt = PromptTemplate(
    template="""Generate a well-structured and professional summary of the given document. The summary should highlight key details, including pricing, categories, discounts, terms, and any other relevant information. Ensure clarity, conciseness, and a formal tone suitable for business or professional use.  Document: {text}\n\n",""",
    input_variables=["text"]
)

def extract_company_info(text: str) -> List[CompanyInfo]:
    """Extracts structured company information from unstructured text."""
    response = (prompt | llm | parser).invoke({"text": text})
    return response.companies if isinstance(response, CompanyInfoList) else []

def extract_data(text: str):
    """Extracts enhanced textual data from unstructured text."""
    response = (enhancement_prompt | llm).invoke({"text": text})
    return response.content

def enhance_extracted_document(text: str):
    """Extracts enhanced textual data from unstructured text."""
    response = (text_enhance_prompt | llm).invoke({"text": text})
    return response.content

def extract_text_from_excel(uploaded_file) -> str:
    """Extracts text from all sheets of an Excel file."""
    wb = load_workbook(filename=uploaded_file, data_only=True)
    extracted_text = ""
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        extracted_text += f"Sheet: {sheet}\n"
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell) if cell is not None else "" for cell in row]
            extracted_text += "\t".join(row_values) + "\n"
        extracted_text += "\n"
    return extracted_text

def extract_text_from_pdf(uploaded_file) -> str:
    """Extracts text from a PDF file."""
    reader = PyPDF2.PdfReader(uploaded_file)
    extracted_text = ""
    for page in reader.pages:
        extracted_text += page.extract_text() + "\n"
    return extracted_text

def extract_text_from_docx(uploaded_file) -> str:
    """Extracts text from a DOCX file, including paragraphs and tables."""
    try:
        doc = DocxDocument(uploaded_file)
        extracted_text = ""
        for para in doc.paragraphs:
            extracted_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    extracted_text += cell.text + "\t"
                extracted_text += "\n"
        return extracted_text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

def ingest_to_pinecone_docs(companies: List[CompanyInfo], embeddings):
    """Ingest approved companies into Pinecone with metadata."""
    st.write("Starting ingestion process for companies...")
    for comp in companies:
        comp_text = (
            f"Name: {comp.name}; Package: {comp.package}; Number of Visas: {comp.number_of_visas}; "
            f"Number of Shareholders: {comp.number_of_shareholders}; Office Required: {comp.office_required}; "
            f"Cost: {comp.cost}; Activity: {comp.activity}"
        )

        metadata = {
            "name": comp.name,
            "package": comp.package,
            "number_of_visas": comp.number_of_visas,
            "number_of_shareholders": comp.number_of_shareholders,
            "office_required": comp.office_required,
            "cost": comp.cost,
            "activity": comp.activity,
            "isMain": False
        }

        document = Document(
            page_content=comp_text,
            metadata=metadata
        )

        if document:
            try:
                PineconeVectorStore.from_documents(
                    documents=[document],
                    embedding=embeddings,
                    index_name=index
                )
                st.success(f"Successfully ingested {comp.name} into Pinecone.")
            except Exception as e:
                st.error(f"Error while ingesting to Pinecone: {e}")
        else:
            st.warning("No document to ingest.")
    return "Approved companies have been ingested into Pinecone."

def ingest_text_to_pinecone(text: str, embeddings):
    """Ingest raw text into Pinecone as a single document."""
    if not text.strip():
        st.error("No text extracted from the uploaded files.")
        return
    document = Document(
        page_content=text,
        metadata={"isMain": True}
    )
    try:
        PineconeVectorStore.from_documents(
            documents=[document],
            embedding=embeddings,
            index_name=index
        )
        st.success("Successfully ingested text into Pinecone.")
    except Exception as e:
        st.error(f"Error while ingesting text: {e}")

# Streamlit App Layout with Tabs
st.title("FZcompare Admin")

try:
    authenticator.login(location='main')
except Exception as e:
    st.error(f"Login widget error: {e}")

# Retrieve authentication status
auth_status = st.session_state.get("authentication_status")

if auth_status is True:
    st.write(f'Welcome **{st.session_state["name"]}**')
    
    tab1, tab2 = st.tabs(["Company Info", "Text Ingestion"])

    # Tab 1: Company Info Extraction & Ingestion
    with tab1:
        st.write("Upload Excel, PDF, or DOCX files to extract, edit, approve, and ingest company information into DB.")

        if "uploaded_files" not in st.session_state:
            st.session_state.uploaded_files = None

        uploaded_files = st.file_uploader(
            "Upload Excel, PDF, or DOCX files", 
            type=["xlsx", "pdf", "docx"], 
            accept_multiple_files=True, 
            key="company_info_upload"
        )
        if uploaded_files:
            st.session_state.uploaded_files = uploaded_files

        if st.button("Extract Company Info"):
            if not st.session_state.uploaded_files:
                st.error("Please upload at least one file.")
            else:
                combined_text = ""
                for uploaded_file in st.session_state.uploaded_files:
                    ext = uploaded_file.name.split('.')[-1].lower()
                    if ext == "xlsx":
                        combined_text += extract_text_from_excel(uploaded_file) + "\n"
                    elif ext == "pdf":
                        combined_text += extract_text_from_pdf(uploaded_file) + "\n"
                    elif ext == "docx":
                        combined_text += extract_text_from_docx(uploaded_file) + "\n"
                    else:
                        st.warning(f"Unsupported file type: {uploaded_file.name}")

                st.info("Enhancing text...")
                enhanced_text = extract_data(combined_text)
                st.session_state.enhanced_text = enhanced_text

                st.info("Extracting Company Information...")
                company_infos = extract_company_info(enhanced_text)
                st.session_state.company_infos = company_infos

        if "company_infos" in st.session_state:
            st.subheader("Extracted Company Information (Editable & Approve Each)")
            approved_companies = []
            for idx, info in enumerate(st.session_state.company_infos):
                with st.expander(f"Company {idx+1} Details", expanded=True):
                    name = st.text_input("Name", value=info.name, key=f"name_{idx}")
                    package = st.text_input("Package", value=info.package, key=f"package_{idx}")
                    number_of_visas = st.number_input("Number of Visas", value=info.number_of_visas or 0, step=1, key=f"visas_{idx}")
                    number_of_shareholders = st.number_input("Number of Shareholders", value=info.number_of_shareholders or 0, step=1, key=f"shareholders_{idx}")
                    office_required = st.checkbox("Office Required", value=info.office_required if info.office_required is not None else False, key=f"office_{idx}")
                    cost = st.number_input("Cost", value=info.cost or 0.0, step=100.0, format="%.2f", key=f"cost_{idx}")
                    activity = st.text_input("Activity", value=info.activity if info.activity is not None else "", key=f"activity_{idx}")
                    
                    approved = st.checkbox("Approve this company", key=f"approve_{idx}")
                    
                    if approved:
                        updated_company = CompanyInfo(
                            name=name,
                            package=package,
                            number_of_visas=number_of_visas,
                            number_of_shareholders=number_of_shareholders,
                            office_required=office_required,
                            cost=cost,
                            activity=activity,
                        )
                        approved_companies.append(updated_company)
            
            if st.button("Ingest Approved Data into Pinecone"):
                if not approved_companies:
                    st.error("No companies approved for ingestion.")
                else:
                    ingest_to_pinecone_docs(approved_companies, embeddings)

    # Tab 2: Raw Text Ingestion
    with tab2:
        st.write("Upload Excel, PDF, or DOCX files to extract and ingest the raw text directly into Pinecone.")
        if "text_upload_files" not in st.session_state:
            st.session_state.text_upload_files = None

        text_uploaded_files = st.file_uploader(
            "Upload Excel, PDF, or DOCX files for text ingestion", 
            type=["xlsx", "pdf", "docx"], 
            accept_multiple_files=True, 
            key="text_ingest_upload"
        )
        if text_uploaded_files:
            st.session_state.text_upload_files = text_uploaded_files

        if st.button("Extract & Ingest Text"):
            if not st.session_state.text_upload_files:
                st.error("Please upload at least one file for text ingestion.")
            else:
                raw_text = ""
                for uploaded_file in st.session_state.text_upload_files:
                    ext = uploaded_file.name.split('.')[-1].lower()
                    if ext == "xlsx":
                        raw_text += extract_text_from_excel(uploaded_file) + "\n"
                    elif ext == "pdf":
                        raw_text += extract_text_from_pdf(uploaded_file) + "\n"
                    elif ext == "docx":
                        raw_text += extract_text_from_docx(uploaded_file) + "\n"
                    else:
                        st.warning(f"Unsupported file type: {uploaded_file.name}")
                enhanced_text = enhance_extracted_document(raw_text)
                st.session_state.enhanced_text = enhanced_text
                st.info("Ingesting extracted text into Pinecone...")
                ingest_text_to_pinecone(raw_text, embeddings)
elif st.session_state.get('authentication_status') is False:
    st.error('Username/password is incorrect')
elif st.session_state.get('authentication_status') is None:
    st.warning('Please enter your username and password')
